"""
app/services/file_service.py
Handles file validation and text extraction for uploaded case documents.
Supports PDF (PyPDF2), DOCX (python-docx), and plain text.
"""
import hashlib
import io
import re
from pathlib import Path

import PyPDF2
import docx
from fastapi import UploadFile

from app.core.config import get_settings
from app.core.exceptions import (
    FileExtractionException,
    FileTooLargeException,
    UnsupportedFileTypeException,
)
from app.core.logger import logger


settings = get_settings()


class FileExtractionResult:
    __slots__ = ("filename", "extension", "size_bytes", "raw_text", "char_count", "content_hash")

    def __init__(
        self,
        filename: str,
        extension: str,
        size_bytes: int,
        raw_text: str,
    ) -> None:
        self.filename     = filename
        self.extension    = extension
        self.size_bytes   = size_bytes
        self.raw_text     = raw_text[: settings.MAX_TEXT_CHARS]
        self.char_count   = len(self.raw_text)
        self.content_hash = hashlib.sha256(self.raw_text.encode()).hexdigest()


class FileService:
    """Validates uploaded files and extracts plain text from them."""

    EXTRACTORS = {
        ".pdf":  "_extract_pdf",
        ".docx": "_extract_docx",
        ".doc":  "_extract_docx",
        ".txt":  "_extract_text",
        ".rtf":  "_extract_text",
    }

    async def validate_and_extract(self, upload: UploadFile) -> FileExtractionResult:
        """
        Full pipeline: validate → read → extract → return result.
        Raises typed exceptions on any failure.
        """
        filename  = upload.filename or "unknown"
        extension = Path(filename).suffix.lower()

        # 1. Validate file type
        if extension not in settings.ALLOWED_EXTENSIONS:
            raise UnsupportedFileTypeException(extension, settings.ALLOWED_EXTENSIONS)

        # 2. Read content into memory
        content = await upload.read()
        size_bytes = len(content)

        # 3. Validate size
        if size_bytes > settings.max_upload_bytes:
            size_mb = size_bytes / (1024 * 1024)
            raise FileTooLargeException(size_mb, settings.MAX_UPLOAD_SIZE_MB)

        if size_bytes == 0:
            raise FileExtractionException(filename, "File is empty.")

        logger.info(
            "File received",
            filename=filename,
            extension=extension,
            size_bytes=size_bytes,
        )

        # 4. Extract text
        extractor_name = self.EXTRACTORS[extension]
        extractor      = getattr(self, extractor_name)
        raw_text       = extractor(content, filename)

        if not raw_text or len(raw_text.strip()) < 20:
            raise FileExtractionException(filename, "Extracted text is too short or empty.")

        logger.info(
            "Text extracted",
            filename=filename,
            char_count=len(raw_text),
            capped=len(raw_text) > settings.MAX_TEXT_CHARS,
        )

        return FileExtractionResult(
            filename=filename,
            extension=extension,
            size_bytes=size_bytes,
            raw_text=raw_text,
        )

    # ── Extractors ───────────────────────────────────────────

    def _extract_pdf(self, content: bytes, filename: str) -> str:
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            if reader.is_encrypted:
                raise FileExtractionException(filename, "PDF is encrypted/password protected.")

            texts: list[str] = []
            max_pages = min(len(reader.pages), 40)

            for i in range(max_pages):
                page = reader.pages[i]
                page_text = page.extract_text() or ""
                texts.append(page_text)

            full_text = "\n".join(texts)
            return self._clean_text(full_text)
        except FileExtractionException:
            raise
        except Exception as e:
            raise FileExtractionException(filename, f"PDF parse error: {e}") from e

    def _extract_docx(self, content: bytes, filename: str) -> str:
        try:
            doc = docx.Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)

            full_text = "\n".join(paragraphs)
            return self._clean_text(full_text)
        except Exception as e:
            raise FileExtractionException(filename, f"DOCX parse error: {e}") from e

    def _extract_text(self, content: bytes, filename: str) -> str:
        try:
            # Try UTF-8 first, fall back to latin-1
            try:
                text = content.decode("utf-8")
            except UnicodeDecodeError:
                text = content.decode("latin-1", errors="replace")
            return self._clean_text(text)
        except Exception as e:
            raise FileExtractionException(filename, f"Text read error: {e}") from e

    # ── Helpers ──────────────────────────────────────────────

    @staticmethod
    def _clean_text(text: str) -> str:
        """Normalise whitespace, remove null bytes and control characters."""
        # Remove null bytes
        text = text.replace("\x00", "")
        # Replace control chars except newlines/tabs
        text = re.sub(r"[\x01-\x08\x0b\x0c\x0e-\x1f\x7f]", " ", text)
        # Collapse excessive whitespace
        text = re.sub(r" {4,}", " ", text)
        text = re.sub(r"\n{4,}", "\n\n\n", text)
        return text.strip()
